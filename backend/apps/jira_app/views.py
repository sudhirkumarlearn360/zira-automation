from django.shortcuts import render, redirect
from django.views import View
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.utils.decorators import method_decorator
import json

from .services.jira_service import JiraService
from .services.ai_service import AIService
from django.views.generic import TemplateView

class HomeView(TemplateView):
    template_name = 'home.html'

class EpicStoriesView(View):
    def get(self, request, epic_key):
        jira_service = JiraService()
        stories = jira_service.get_stories_for_epic(epic_key)
        
        # Base URL for links
        import os
        jira_base_url = os.environ.get('JIRA_URL', '')
        
        context = {
            'epic_key': epic_key,
            'stories': stories,
            'jira_base_url': jira_base_url
        }
        return render(request, 'jira/epic_stories.html', context)

class IndexView(View):
    def get(self, request):
        return render(request, 'jira/index.html')

class StoryDetailView(View):
    def get(self, request, story_key):
        jira_service = JiraService()
        story = jira_service.get_story_details(story_key)
        return render(request, 'jira/story_detail.html', {'story': story})


@method_decorator(csrf_exempt, name='dispatch')
class PreviewTasksView(View):
    """
    Step 1: Receive story keys & task type -> Return AI generated JSON for review.
    """
    def post(self, request):
        try:
            data = json.loads(request.body)
            story_keys = data.get('story_keys', [])
            epic_key = data.get('epic_key', 'UNKNOWN')
            task_type = data.get('task_type', 'General')
            
            # Re-fetch stories to get text (in real app, use cache/db)
            jira_service = JiraService()
            ai_service = AIService()
            all_stories = jira_service.get_stories_for_epic(epic_key)
            story_map = {s['key']: s for s in all_stories}
            
            previews = []
            
            for key in story_keys:
                story_data = story_map.get(key)
                
                # Fallback: If not finding story in bulk fetch (e.g. single story view context), fetch it directly
                if not story_data:
                    try:
                        detail = jira_service.get_story_details(key)
                        if detail:
                            story_data = detail
                    except Exception as e:
                        print(f"Failed to fetch individual story {key}: {e}")

                if not story_data:
                    print(f"Skipping {key} - not found.")
                    continue
                
                full_text = f"Summary: {story_data['summary']}\nDescription: {story_data['description']}"
                
                # Generate AI Content
                ai_json = ai_service.generate_task_from_story(full_text, task_type)
                
                if ai_json:
                    previews.append({
                        'original_key': key,
                        'generated_content': ai_json
                    })
            
            return JsonResponse({'previews': previews})
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)

@method_decorator(csrf_exempt, name='dispatch')
class CreateConfirmedTasksView(View):
    """
    Step 2: Receive CONFIRMED JSON content -> Create in Jira.
    """
    def post(self, request):
        try:
            data = json.loads(request.body)
            approved_tasks = data.get('tasks', []) # List of {original_key, generated_content}
            epic_key = data.get('epic_key')
            
            jira_service = JiraService()
            results = []
            
            for task_item in approved_tasks:
                original_key = task_item.get('original_key')
                content = task_item.get('generated_content')
                
                # Create Task
                new_task = jira_service.create_task(content, parent_key=epic_key)
                
                if new_task:
                    results.append({'key': original_key, 'status': 'CREATED', 'new_task_key': new_task.get('key')})
                else:
                    results.append({'key': original_key, 'status': 'FAILED', 'reason': 'Jira Create Error'})

            return JsonResponse({'results': results})
            
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)


# ==================== STORIES GENERATOR VIEWS ====================

class StoriesView(View):
    """Main UI view for /jira/stories/ endpoint."""
    def get(self, request):
        return render(request, 'jira/stories.html')


@method_decorator(csrf_exempt, name='dispatch')
class PreviewStoryTaskAPIView(View):
    """API endpoint to generate AI task payload from a story (Preview mode)."""
    def post(self, request):
        try:
            data = json.loads(request.body)
            story_key = data.get('story_key')
            task_type = data.get('task_type', 'General')
            additional_context = data.get('additional_context', '')

            if not story_key:
                return JsonResponse({'error': 'Story key is required'}, status=400)

            jira_service = JiraService()
            ai_service = AIService()

            # Fetch story details
            story_detail = jira_service.get_story_details(story_key)
            if not story_detail:
                return JsonResponse({'error': f'Failed to fetch story {story_key}. It may not exist or require authentication.'}, status=404)

            # Build content payload for AI
            content_for_ai = f"Summary: {story_detail['summary']}\nDescription: {story_detail.get('description', '')}"

            # Request generation
            ai_json = ai_service.generate_task_from_story(content_for_ai, task_type, additional_context)
            if not ai_json:
                print(f"Error: AI generation failed for story {story_key}")
                return JsonResponse({'error': 'Failed to generate task from AI service.'}, status=500)

            return JsonResponse({
                'success': True,
                'message': 'AI Task Preview Generated',
                'story_key': story_key,
                'generated_content': ai_json
            })

        except Exception as e:
            print(f"Exception in PreviewStoryTaskAPIView: {e}")
            return JsonResponse({'error': str(e)}, status=500)


@method_decorator(csrf_exempt, name='dispatch')
class CreateStoryTaskAPIView(View):
    """API endpoint to create a Jira task from a confirmed AI payload."""
    def post(self, request):
        try:
            data = json.loads(request.body)
            story_key = data.get('story_key')
            ai_payload = data.get('ai_payload')

            if not story_key or not ai_payload:
                return JsonResponse({'error': 'Story key and AI payload are required'}, status=400)

            jira_service = JiraService()

            # Create in Jira
            print(f"Attempting to create task with edited AI JSON: {ai_payload}")
            new_task = jira_service.create_task(ai_payload, parent_key=story_key)

            if new_task and 'key' in new_task:
                 return JsonResponse({
                     'success': True,
                     'message': 'Task Successfully Created',
                     'story_key': story_key,
                     'new_task_key': new_task['key']
                 })
            else:
                 print(f"Error: Jira create_task returned {new_task}")
                 return JsonResponse({'error': 'Failed to create task in Jira.', 'details': new_task}, status=500)

        except Exception as e:
            print(f"Exception in CreateStoryTaskAPIView: {e}")
            return JsonResponse({'error': str(e)}, status=500)


# ==================== TEST CASE GENERATOR VIEWS ====================

class TestCaseGeneratorView(View):
    """Main page for Test Case Generator"""
    def get(self, request):
        import os
        jira_base_url = os.environ.get('JIRA_URL', '')
        return render(request, 'jira/test_case_generator.html', {
            'jira_base_url': jira_base_url
        })


@method_decorator(csrf_exempt, name='dispatch')
class GenerateTestCasesView(View):
    """Generate test cases from story/epic/task content"""
    def post(self, request):
        try:
            data = json.loads(request.body)
            source_type = data.get('source_type', 'story')  # epic, story, task
            source_key = data.get('source_key', '')
            user_prompt = data.get('prompt', '')
            use_categorized = data.get('categorized', True)  # Default to categorized for epic
            
            jira_service = JiraService()
            ai_service = AIService()
            
            # Fetch content based on source type
            content = ""
            source_data = {}
            
            if source_type == 'epic':
                # Fetch all stories for the epic
                stories = jira_service.get_stories_for_epic(source_key)
                stories_for_ai = [
                    {'key': s['key'], 'summary': s['summary'], 'description': s.get('description', '')}
                    for s in stories
                ]
                source_data = {'epic_key': source_key, 'story_count': len(stories), 'stories': stories}
                
                if use_categorized:
                    # Use categorized generation for Epic
                    result = ai_service.generate_categorized_epic_test_cases(stories_for_ai, user_prompt)
                    
                    return JsonResponse({
                        'success': True,
                        'source_type': source_type,
                        'source_key': source_key,
                        'source_data': source_data,
                        'categorized': True,
                        'categories': result.get('categories', {}),
                        'summary': result.get('summary', {}),
                        'original_stories': stories_for_ai
                    })
                else:
                    # Flat generation
                    content_parts = []
                    for story in stories:
                        content_parts.append(f"Story {story['key']}: {story['summary']}\n{story.get('description', '')}")
                    content = "\n\n---\n\n".join(content_parts)
                    result = ai_service.generate_test_cases(content, source_type, user_prompt)
            else:
                # Single story or task
                story = jira_service.get_story_details(source_key)
                if story:
                    content = f"Summary: {story['summary']}\nDescription: {story.get('description', '')}"
                    source_data = story
                else:
                    return JsonResponse({'error': f'Could not fetch {source_type}: {source_key}'}, status=404)
                
                result = ai_service.generate_test_cases(content, source_type, user_prompt)
            
            return JsonResponse({
                'success': True,
                'source_type': source_type,
                'source_key': source_key,
                'source_data': source_data,
                'categorized': False,
                'test_cases': result.get('test_cases', []),
                'summary': result.get('summary', ''),
                'original_content': content
            })
            
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)


@method_decorator(csrf_exempt, name='dispatch')
class RefineTestCasesView(View):
    """Refine existing test cases with additional prompt"""
    def post(self, request):
        try:
            data = json.loads(request.body)
            current_test_cases = data.get('test_cases', [])
            refinement_prompt = data.get('prompt', '')
            original_content = data.get('original_content', '')
            
            ai_service = AIService()
            result = ai_service.refine_test_cases(current_test_cases, refinement_prompt, original_content)
            
            return JsonResponse({
                'success': True,
                'test_cases': result.get('test_cases', []),
                'changes_made': result.get('changes_made', [])
            })
            
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)


@method_decorator(csrf_exempt, name='dispatch')
class RegenerateCategoryTestCasesView(View):
    """Regenerate test cases for a specific category (BE/FE/QA/OTHER)"""
    def post(self, request):
        try:
            data = json.loads(request.body)
            category = data.get('category', '')  # BE, FE, QA, OTHER
            stories = data.get('stories', [])
            current_test_cases = data.get('test_cases', [])
            user_prompt = data.get('prompt', '')
            original_context = data.get('original_context', '')
            
            if not category:
                return JsonResponse({'error': 'Category is required'}, status=400)
            
            ai_service = AIService()
            result = ai_service.regenerate_category_test_cases(
                category, stories, current_test_cases, user_prompt, original_context
            )
            
            return JsonResponse({
                'success': True,
                'category': category,
                'test_cases': result.get('test_cases', []),
                'changes_made': result.get('changes_made', [])
            })
            
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)


@method_decorator(csrf_exempt, name='dispatch')
class ExportTestCasesView(View):
    """Export test cases to Excel or Word format"""
    def post(self, request):
        from django.http import HttpResponse
        import io
        import traceback
        
        try:
            data = json.loads(request.body)
            test_cases = data.get('test_cases', [])
            export_format = data.get('format', 'excel')  # excel or word
            source_info = data.get('source_info', '')
            
            print(f"[EXPORT] Format: {export_format}, Test Cases: {len(test_cases)}")
            
            if export_format == 'excel':
                return self._export_excel(test_cases, source_info)
            else:
                return self._export_word(test_cases, source_info)
                
        except Exception as e:
            print(f"[EXPORT ERROR] {str(e)}")
            print(traceback.format_exc())
            return JsonResponse({'error': str(e)}, status=500)
    
    def _export_excel(self, test_cases, source_info):
        from django.http import HttpResponse
        import io
        import traceback
        
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        except ImportError:
            return JsonResponse({'error': 'openpyxl not installed. Run: pip install openpyxl'}, status=500)
        
        try:
            wb = Workbook()
            ws = wb.active
            ws.title = "Test Cases"
            
            # Header styling
            header_font = Font(bold=True, color="FFFFFF")
            header_fill = PatternFill(start_color="0052CC", end_color="0052CC", fill_type="solid")
            thin_border = Border(
                left=Side(style='thin'),
                right=Side(style='thin'),
                top=Side(style='thin'),
                bottom=Side(style='thin')
            )
            
            # Headers
            headers = ['ID', 'Title', 'Description', 'Preconditions', 'Steps', 'Expected Result', 'Priority', 'Type']
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=1, column=col, value=header)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = Alignment(horizontal='center', vertical='center')
                cell.border = thin_border
            
            # Data rows
            for row, tc in enumerate(test_cases, 2):
                # Handle preconditions and steps that might be strings or lists
                preconditions = tc.get('preconditions', [])
                if isinstance(preconditions, list):
                    preconditions_str = '\n'.join(str(p) for p in preconditions)
                else:
                    preconditions_str = str(preconditions) if preconditions else ''
                
                steps = tc.get('steps', [])
                if isinstance(steps, list):
                    steps_str = '\n'.join(str(s) for s in steps)
                else:
                    steps_str = str(steps) if steps else ''
                
                ws.cell(row=row, column=1, value=str(tc.get('id', ''))).border = thin_border
                ws.cell(row=row, column=2, value=str(tc.get('title', ''))).border = thin_border
                ws.cell(row=row, column=3, value=str(tc.get('description', ''))).border = thin_border
                ws.cell(row=row, column=4, value=preconditions_str).border = thin_border
                ws.cell(row=row, column=5, value=steps_str).border = thin_border
                ws.cell(row=row, column=6, value=str(tc.get('expected_result', ''))).border = thin_border
                ws.cell(row=row, column=7, value=str(tc.get('priority', ''))).border = thin_border
                ws.cell(row=row, column=8, value=str(tc.get('type', tc.get('category', '')))).border = thin_border
            
            # Auto-adjust column widths (safely)
            for col in ws.columns:
                max_length = 0
                column = col[0].column_letter
                for cell in col:
                    try:
                        if cell.value and len(str(cell.value)) > max_length:
                            max_length = min(len(str(cell.value)), 100)  # Cap at 100 chars
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                ws.column_dimensions[column].width = adjusted_width
            
            # Save to buffer
            buffer = io.BytesIO()
            wb.save(buffer)
            buffer.seek(0)
            
            response = HttpResponse(
                buffer.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
            )
            response['Content-Disposition'] = 'attachment; filename="test_cases.xlsx"'
            return response
        except Exception as e:
            print(f"[EXCEL EXPORT ERROR] {str(e)}")
            print(traceback.format_exc())
            return JsonResponse({'error': f'Excel export failed: {str(e)}'}, status=500)
    
    def _export_word(self, test_cases, source_info):
        from django.http import HttpResponse
        import io
        
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            return JsonResponse({'error': 'python-docx not installed. Run: pip install python-docx'}, status=500)
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Test Cases Document', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if source_info:
            doc.add_paragraph(f"Source: {source_info}")
        
        doc.add_paragraph(f"Total Test Cases: {len(test_cases)}")
        doc.add_paragraph("")
        
        # Test cases
        for tc in test_cases:
            doc.add_heading(f"{tc.get('id', 'TC')} - {tc.get('title', 'Untitled')}", level=2)
            
            doc.add_paragraph(f"Priority: {tc.get('priority', 'N/A')} | Type: {tc.get('type', 'N/A')}")
            
            doc.add_heading('Description', level=3)
            doc.add_paragraph(tc.get('description', 'No description'))
            
            doc.add_heading('Preconditions', level=3)
            preconditions = tc.get('preconditions', [])
            if isinstance(preconditions, list):
                for pre in preconditions:
                    doc.add_paragraph(f"• {pre}")
            else:
                doc.add_paragraph(f"• {preconditions}" if preconditions else "None")
            
            doc.add_heading('Steps', level=3)
            steps = tc.get('steps', [])
            if isinstance(steps, list):
                for i, step in enumerate(steps, 1):
                    doc.add_paragraph(f"{i}. {step}")
            else:
                doc.add_paragraph(str(steps) if steps else "No steps")
            
            doc.add_heading('Expected Result', level=3)
            doc.add_paragraph(tc.get('expected_result', 'No expected result specified'))
            
            doc.add_paragraph("")  # Spacing between test cases
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename="test_cases.docx"'
        return response

